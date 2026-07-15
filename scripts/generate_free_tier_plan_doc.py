"""Generates Free_Tier_Resource_Plan.docx - the manager-facing summary of the
fully card-free resource stack for Milestone 1.

Run with: .venv/Scripts/python.exe scripts/generate_free_tier_plan_doc.py
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUTPUT_PATH = Path(__file__).resolve().parent.parent / "Free_Tier_Resource_Plan.docx"

ACCENT = RGBColor(0x1F, 0x4E, 0x79)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = ACCENT
    return h


def add_bullets(doc, items, bold_lead=True):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        if bold_lead and ":" in item:
            lead, rest = item.split(":", 1)
            run = p.add_run(lead + ":")
            run.bold = True
            p.add_run(rest)
        else:
            p.add_run(item)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    return table


def build():
    doc = Document()

    title = doc.add_heading("Free-Tier Resource Plan (Revised)", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(
        "Model-Independent AI Engine & Resume Builder \u2014 Milestone 1\n"
        "Revision: fully card-free stack, verified line-by-line"
    )
    run.italic = True

    add_heading(doc, "1. Why This Plan Was Revised", level=1)
    doc.add_paragraph(
        "The original free-tier plan relied on Render and, briefly, Belmo.io for "
        "hosting. Both surfaced problems in practice that contradicted their own "
        "marketing: Render repeatedly prompted for a credit card despite its "
        "documented card-free free tier, and Belmo gated Docker-based deployments "
        "behind a paid Pro plan. Rather than keep chasing platforms whose real-world "
        "behavior doesn't match their stated policy, this revision drops every "
        "component that depends on a GUI-driven cloud signup flow and replaces it "
        "with options that are either already working or verifiably require zero "
        "account and zero card."
    )

    add_heading(doc, "2. Revised Stack \u2014 What's Actually Free, Verified", level=1)
    add_table(
        doc,
        ["Component", "Service", "Card / Account Required?", "Status"],
        [
            ["LLM access", "OpenRouter (free-tier models)", "Free account, no card", "Done \u2014 API key configured"],
            ["Version control & CI", "GitHub + GitHub Actions", "Free account, no card", "Done \u2014 repo live, CI configured"],
            ["Runtime hosting", "Local machine (uvicorn)", "None \u2014 no account at all", "Ready to run"],
            ["Public URL", "Cloudflare Quick Tunnel (cloudflared)", "None \u2014 no account, no card, no signup", "Ready to run"],
            ["Workflow state / checkpoints", "SQLite (local file)", "None", "Already the default"],
            ["Uploaded file storage", "Local disk", "None", "Already the default"],
            ["OCR", "Tesseract (self-hosted)", "None", "Already integrated"],
            ["Containerization", "Docker (kept for later)", "None to build/run locally", "Available, not required now"],
        ],
    )

    add_heading(doc, "3. What Was Dropped and Why", level=1)
    add_bullets(
        doc,
        [
            "Render: dropped \u2014 free tier repeatedly triggered a card-verification "
            "prompt in practice, despite documentation stating no card is required.",
            "Belmo.io: dropped \u2014 Docker-based deploys are gated behind a paid "
            "\"Pro\" plan; free tier does not cover our deployment method.",
            "Neon (Postgres): dropped \u2014 not needed. SQLite already ships as the "
            "default checkpoint backend and needs no external account.",
            "Cloudflare R2 (object storage): dropped \u2014 Cloudflare's own community "
            "forum confirms R2 requires a card on file to activate, even on the free "
            "tier. Local disk storage is already the default and needs nothing.",
        ],
    )

    add_heading(doc, "4. How Public Access Works Without Any Signup", level=1)
    doc.add_paragraph(
        "Cloudflare's \"Quick Tunnel\" feature (part of the free cloudflared CLI "
        "tool) creates an instant, secure, HTTPS-enabled public URL "
        "(*.trycloudflare.com) that forwards to the engine running locally \u2014 with "
        "a single command and no Cloudflare account of any kind:"
    )
    code_p = doc.add_paragraph()
    code_run = code_p.add_run("cloudflared tunnel --url http://localhost:8000")
    code_run.font.name = "Consolas"
    code_run.font.size = Pt(10)

    doc.add_paragraph(
        "Trade-off: this URL is temporary \u2014 it changes if the tunnel is "
        "restarted, and it is capped at 200 concurrent requests (far more than "
        "needed for demos and QA). It is not a 24/7 production deployment. That "
        "matches Section 8 of the original scope document, which already states "
        "that cloud hosting decisions are deferred until after the MVP has been "
        "tested \u2014 this plan simply gets a working, shareable demo live now, at "
        "zero cost and zero account risk, without pre-committing to a hosting "
        "provider before the MVP is validated."
    )

    add_heading(doc, "5. Path to a Permanent Deployment (Post-MVP)", level=1)
    doc.add_paragraph(
        "Once the MVP is validated and the client/manager decides on a hosting "
        "budget, moving to a real always-on deployment (AWS/GCP/Azure, or a paid "
        "Render/Railway plan) is a configuration change, not a rebuild \u2014 the "
        "engine already runs in Docker and is model-independent via OpenRouter. "
        "No code changes are required to move from local + tunnel to a managed "
        "cloud host."
    )

    add_heading(doc, "6. Summary", level=1)
    doc.add_paragraph(
        "This revised plan uses only services that have been individually "
        "verified \u2014 not just documented \u2014 to require no credit card and no "
        "paid plan: OpenRouter's free models, GitHub's free tier, and Cloudflare's "
        "account-free Quick Tunnel, combined with the engine's own local-first "
        "defaults (SQLite, local disk, Tesseract). The result is a fully working, "
        "publicly reachable Milestone 1 demo at $0 cost with zero payment "
        "information entered anywhere."
    )

    doc.save(OUTPUT_PATH)
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
