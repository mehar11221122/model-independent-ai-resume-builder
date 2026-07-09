"""One-off script to generate Free_Tier_Resource_Plan.docx.

Not part of the running application - just a helper to produce a shareable
Word document summarizing free/low-cost infrastructure options for the
project, for non-technical stakeholders (e.g. a manager).
"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = "Free_Tier_Resource_Plan.docx"

ACCENT = RGBColor(0x1F, 0x4E, 0x79)


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        set_cell_shading(hdr_cells[i], "1F4E79")

    for row_data in rows:
        row_cells = table.add_row().cells
        for i, value in enumerate(row_data):
            row_cells[i].text = ""
            p = row_cells[i].paragraphs[0]
            run = p.add_run(value)
            run.font.size = Pt(10)

    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Inches(width)

    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = ACCENT
    return heading


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_document():
    doc = Document()

    # --- Title block -----------------------------------------------------
    title = doc.add_heading("Free & Low-Cost Resource Plan", level=0)
    for run in title.runs:
        run.font.color.rgb = ACCENT
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = subtitle.add_run(
        "Model-Independent AI Engine & Resume Builder Module"
    )
    run.bold = True
    run.font.size = Pt(13)

    meta = doc.add_paragraph()
    meta.add_run("Prepared for: Management  |  Version 1.0  |  July 2026").italic = True

    doc.add_paragraph()

    # --- Purpose -----------------------------------------------------------
    add_heading(doc, "1. Purpose", level=1)
    doc.add_paragraph(
        "This document lists the free or near-free versions of every third-party "
        "service, model, and piece of infrastructure required to build and run the "
        "AI Engine and Resume Builder module (as defined in the Project Scope "
        "Document), so the project can be developed and demoed at little to no "
        "cost before committing budget to paid tiers."
    )
    doc.add_paragraph(
        "Because the engine is model-independent and every backend is a "
        "configuration choice (never hard-coded), switching from a free "
        "resource to a paid one later is a settings change, not a rebuild."
    )

    # --- Summary table -------------------------------------------------
    add_heading(doc, "2. Summary: Paid Component vs. Free Alternative", level=1)
    add_table(
        doc,
        ["Category", "Paid Option", "Free Alternative", "Key Limit"],
        [
            ("LLM inference", "Claude Sonnet 4 / GPT-4o / Gemini Pro",
             "OpenRouter free models (DeepSeek, Qwen, Llama 3.3, GPT-OSS)",
             "~20 req/min, 200 req/day per model"),
            ("Compute / hosting", "AWS / GCP / Azure paid VM",
             "Render free web service, or Oracle Cloud Always Free VM",
             "Render: spins down after 15 min idle; Oracle: 2 OCPU / 12GB RAM"),
            ("Database (checkpointing)", "Managed Postgres",
             "Neon free tier", "0.5GB storage, 100 compute-hrs/month, permanent"),
            ("Cache / queue", "Managed Redis",
             "Upstash free tier", "256MB data, 500K commands/month"),
            ("File storage", "AWS S3",
             "Cloudflare R2 free tier", "10GB storage, zero egress fees"),
            ("OCR (scanned documents)", "Google Vision / AWS Textract",
             "Tesseract OCR (self-hosted)", "No hard limit; self-managed"),
            ("Observability / tracing", "LangSmith paid tier",
             "LangSmith free tier or self-hosted Langfuse",
             "~5K traces/month (LangSmith free)"),
            ("CI/CD", "Paid CI minutes",
             "GitHub Actions", "2,000 min/month free (private repos)"),
        ],
        col_widths=[1.3, 1.7, 2.3, 1.8],
    )

    # --- Detailed sections -----------------------------------------------
    add_heading(doc, "3. Detail by Category", level=1)

    add_heading(doc, "3.1 LLM Inference (Model Gateway)", level=2)
    doc.add_paragraph(
        "OpenRouter itself has no subscription fee — cost only comes from "
        "per-token usage of paid models. Models with a \":free\" suffix "
        "(e.g. deepseek/deepseek-chat-v3-0324:free, "
        "meta-llama/llama-3.3-70b-instruct:free, openai/gpt-oss-120b:free) "
        "cost nothing at all and require no credit card."
    )
    add_bullets(doc, [
        "Trade-off: free models are rate-limited (~20 requests/minute, 200/day) "
        "and generally lower quality than Claude Sonnet 4, especially for "
        "nuanced Arabic generation.",
        "Recommended path: use free models through development and demos; "
        "budget for Claude Sonnet 4 (or another paid model) once quality "
        "needs to be production-grade, particularly for Arabic resumes.",
    ])

    add_heading(doc, "3.2 Compute / Hosting", level=2)
    add_bullets(doc, [
        "Render (free): git-push deployment, 750 free instance-hours/month, "
        "free managed HTTPS. Spins down after 15 minutes of inactivity "
        "(~1 minute cold start on the next request). No persistent disk on "
        "the free tier, so external database/storage is required (see below).",
        "Oracle Cloud Always Free: 2 OCPU + 12GB RAM VM, persistent block "
        "storage, no spin-down, free forever. Requires manual server setup "
        "(SSH, Docker) rather than one-click deploy.",
        "Google Cloud Run: 2 million requests/month free, scales to zero.",
        "Fly.io: no longer offers a free tier as of 2024 — new signups only "
        "get a 2-hour/7-day trial, then require a paid card. Not recommended.",
    ])

    add_heading(doc, "3.3 Database (Workflow State / Checkpointing)", level=2)
    add_bullets(doc, [
        "Neon (recommended): permanent free Postgres — 0.5GB storage, "
        "100 compute-hours/month, scales to zero, no credit card, no "
        "expiration date.",
        "Render Postgres (free): 1GB storage, but expires 30 days after "
        "creation with a 14-day grace period — not suitable as a long-term "
        "free option.",
    ])

    add_heading(doc, "3.4 File Storage (Uploaded Documents)", level=2)
    add_bullets(doc, [
        "Cloudflare R2 (recommended): 10GB free storage, zero egress fees, "
        "S3-compatible API — works as a drop-in replacement for AWS S3.",
        "Local disk storage: zero cost if hosting on a VM with persistent "
        "storage (e.g. Oracle Always Free); not viable on Render's free tier "
        "since it has no persistent disk.",
    ])

    add_heading(doc, "3.5 OCR (Scanned Documents / Images)", level=2)
    doc.add_paragraph(
        "Tesseract OCR is free, open-source, and self-hosted — already the "
        "default in the engine. No per-page cost or quota. Google Cloud "
        "Vision offers a free tier of 1,000 units/month if cloud OCR is "
        "preferred over self-hosting."
    )

    add_heading(doc, "3.6 Observability", level=2)
    doc.add_paragraph(
        "LangSmith is optional; its free Developer tier covers roughly "
        "5,000 traces/month. Langfuse is a free, open-source, self-hosted "
        "alternative with no volume limit if that free tier is outgrown."
    )

    # --- Recommendation -----------------------------------------------
    add_heading(doc, "4. Recommended Zero-Cost Stack for Development & Demo", level=1)
    doc.add_paragraph(
        "The following combination allows the entire engine to run at "
        "$0/month, with no credit card required for any component:"
    )
    add_table(
        doc,
        ["Layer", "Choice"],
        [
            ("Compute", "Render (free web service)"),
            ("Database", "Neon (free Postgres)"),
            ("File storage", "Cloudflare R2 (free tier)"),
            ("LLM inference", "OpenRouter free models"),
            ("OCR", "Tesseract (self-hosted)"),
            ("CI/CD", "GitHub Actions (free tier)"),
        ],
        col_widths=[2.0, 4.0],
    )
    doc.add_paragraph(
        "This stack is sufficient for development, internal testing, and "
        "stakeholder demos. It is not recommended for production traffic "
        "due to rate limits on free LLM models and cold-start delays on "
        "free compute."
    )

    # --- Next steps -----------------------------------------------
    add_heading(doc, "5. Next Steps / Decisions Needed", level=1)
    add_bullets(doc, [
        "Approve using the free-tier stack above through the development "
        "and demo phase of Milestone 1.",
        "Decide the point at which to upgrade to paid LLM models (Claude "
        "Sonnet 4) for production-quality output, especially Arabic.",
        "Decide the point at which to upgrade compute/hosting from free "
        "tiers to a paid, production-grade environment.",
    ])

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run(
        "Note: free-tier limits and availability are accurate as of July 2026 "
        "and are subject to change by each provider."
    ).italic = True

    doc.save(OUTPUT_PATH)
    print(f"Saved to {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
