import io

import docx

from app.ingestion.models import ExtractedDocument, InputKind


def load_docx(source_name: str, data: bytes) -> ExtractedDocument:
    document = docx.Document(io.BytesIO(data))

    parts: list[str] = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                parts.append(row_text)

    text = "\n".join(parts).strip()
    warnings = [] if text else ["No extractable text found in the Word document."]

    return ExtractedDocument(
        source_name=source_name, kind=InputKind.DOCX, raw_text=text, warnings=warnings
    )
